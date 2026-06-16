"""Generate resume v2 — human-sounding, ATS-clean, with REAL clickable hyperlinks
(email, LinkedIn, GitHub). Classic black single-column layout (no template-y
accent colors). Run: ./.venv/bin/python build_resume_v2.py
Output: data/resume/Michael_Yisa_Resume_2026_v2.docx"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import RESUME_DIR

LINK_BLUE = "1A0DAB"


def add_hyperlink(paragraph, url, text, size=9.5):
    """Insert a real external hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), LINK_BLUE); rpr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run)
    paragraph._p.append(link)


def doc():
    d = Document()
    for s in d.sections:
        s.top_margin = s.bottom_margin = Pt(40)
        s.left_margin = s.right_margin = Pt(54)
    st = d.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)
    return d


def heading(d, text):
    """Section heading with a thin bottom rule — standard human resume look."""
    p = d.add_paragraph(); p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(10.5)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in [("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "888888")]:
        bottom.set(qn(k), v)
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def bullet(d, text):
    p = d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    p.add_run(text)


def role(d, title, dates):
    p = d.add_paragraph(); p.paragraph_format.space_before = Pt(4)
    r = p.add_run(title); r.bold = True
    tab = p.add_run("\t" + dates); tab.italic = True; tab.font.size = Pt(9.5)


d = doc()

# header
name = d.add_paragraph(); name.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = name.add_run("Michael Yisa"); r.bold = True; r.font.size = Pt(19)
sub = d.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Full-Stack Software Engineer  ·  AWS Certified Solutions Architect")
r.font.size = Pt(10.5)
c = d.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
c.add_run("Austin, TX  ·  ").font.size = Pt(9.5)
add_hyperlink(c, "mailto:michaelyisa2024@gmail.com", "michaelyisa2024@gmail.com")
c.add_run("  ·  ").font.size = Pt(9.5)
add_hyperlink(c, "https://www.linkedin.com/in/michael-yisa-382a9b249", "LinkedIn")
c.add_run("  ·  ").font.size = Pt(9.5)
add_hyperlink(c, "https://github.com/mikeyisa", "GitHub")
c.add_run("  ·  U.S. Citizen, Security Clearance Eligible").font.size = Pt(9.5)

# summary — plainer, human voice
heading(d, "Summary")
d.add_paragraph(
    "Full-stack engineer with production experience building REST APIs, B2B "
    "integrations, and AWS serverless data pipelines in C#/.NET, Angular, and SQL "
    "Server. AWS Solutions Architect certified, with hands-on Lambda, Step Functions, "
    "and S3 work and a couple of applied AI projects using the Claude API. A U.S. "
    "citizen and clearance-eligible, I'm looking for cloud and AI software roles in "
    "defense and govtech."
).paragraph_format.space_after = Pt(2)

heading(d, "Certifications")
bullet(d, "AWS Certified Solutions Architect – Associate (SAA-C03), 2026")
bullet(d, "Google Advanced Data Analytics; Google Machine Learning; Anthropic Claude Code; "
          "Google Analytics 4")

heading(d, "Experience")
p = d.add_paragraph(); p.add_run("SHI International Corp — Austin, TX").bold = True
role(d, "Associate Software Engineer (Full-Stack, .NET)", "Dec 2025 – Present")
for b in [
    "Build and maintain a cXML B2B order-ingestion API that loads customer purchase "
    "orders into SQL Server, with credential lookups across joined tables in LINQ "
    "surfaced through an Angular front end.",
    "Keep an AWS serverless ETL pipeline (Lambda, Step Functions, S3, Elastic Beanstalk) "
    "running for supplier product feeds landing in SQL Server and MySQL.",
    "Cut product-image latency by about half by serving images from S3 through a "
    "CloudFront distribution.",
    "Dropped compute time roughly 60% by batching downstream calls through an SQS queue "
    "instead of processing them one at a time.",
    "Added new product attributes (end-of-life, ship date) from supplier mappings all "
    "the way through the ETL to the customer-facing product API.",
    "Ship to production through Jenkins with SonarQube checks; cleared a blocked release "
    "by resolving NuGet vulnerabilities (OpenTelemetry pinning on .NET 6).",
    "Write xUnit tests and refactor older code toward repository patterns and dependency "
    "injection; debug across C#/.NET, Angular, React, and SQL Server.",
]:
    bullet(d, b)
role(d, "Jr. IT Analyst Intern", "May 2025 – Dec 2025")
bullet(d, "Built a customer-segmentation dashboard over 1,000+ accounts that surfaced a "
          "$250K+ pipeline opportunity from inactive and high-growth segments.")
bullet(d, "Trained a CatBoost churn model on RFM features and wired the dashboard logic "
          "into Azure Data Lake for dynamic reporting.")

heading(d, "Projects")
p = d.add_paragraph(); r = p.add_run("JobFinder — Agentic Job-Search Platform"); r.bold = True
p.add_run("   ("); add_hyperlink(p, "https://github.com/mikeyisa/jobfinder-yisa-2026", "github.com/mikeyisa/jobfinder-yisa-2026", size=10)
p.add_run(")")
bullet(d, "Built a multi-agent system in FastAPI and the Claude API that pulls live job "
          "postings, scores each against my resume, rewrites the resume for the role, and "
          "drafts recruiter outreach.")
bullet(d, "Handled the agent orchestration, prompt caching to keep costs down, and a "
          "single Docker container running both the API and UI for AWS deployment.")
p = d.add_paragraph(); p.add_run("AI Trading Copilot — Multi-Agent Market Analysis").bold = True
bullet(d, "Built a multi-agent system (Next.js, FastAPI, PostgreSQL, Claude API) that pulls "
          "live market data and runs order-flow, backtesting, and prediction agents to "
          "generate directional analysis.")

heading(d, "Technical Skills")
bullet(d, "Languages: C#, Python, SQL, TypeScript/JavaScript, DAX")
bullet(d, "Frameworks & APIs: .NET 6/8, Angular, React, REST APIs, FastAPI, Next.js, xUnit, "
          "LINQ / Entity Framework")
bullet(d, "Cloud & DevOps: AWS (Lambda, Step Functions, S3, CloudFront, SQS, Elastic "
          "Beanstalk), Azure Data Lake, Docker, Jenkins, SonarQube, Git")
bullet(d, "Data & AI: SQL Server, MySQL, PostgreSQL, Pandas, Scikit-Learn, CatBoost, "
          "Power BI, Claude API and multi-agent systems")

heading(d, "Education")
d.add_paragraph("B.S. Computer Science, Minor in Applied Mathematics — Texas State "
                "University, Dec 2025").paragraph_format.space_after = Pt(2)

heading(d, "Leadership")
bullet(d, "Vice President, Alpha Phi Alpha Fraternity (3-year member)")
bullet(d, "Vice President, Texas State Data Analytics Club")

out = RESUME_DIR / "Michael_Yisa_Resume_2026_v2.docx"
d.save(out)
print("wrote", out)
