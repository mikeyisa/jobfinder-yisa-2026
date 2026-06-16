"""Build a FULL tailored résumé .docx for a specific job — same layout/identity as
the master v2, but with the Claude-tailored headline, summary, and experience
bullets dropped in. Truthful: only the real experience is used, reworded/reordered
for the role. Returns the saved file path."""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import RESUME_DIR

TAILORED_DIR = RESUME_DIR / "tailored"
TAILORED_DIR.mkdir(parents=True, exist_ok=True)
LINK_BLUE = "1A0DAB"

# ---- static, real content (everything except the tailored top + SHI bullets) ----
JR_BULLETS = [
    "Built a customer-segmentation dashboard over 1,000+ accounts that surfaced a "
    "$250K+ pipeline opportunity from inactive and high-growth segments.",
    "Trained a CatBoost churn model on RFM features and wired the dashboard logic "
    "into Azure Data Lake for dynamic reporting.",
]
PROJECTS = [
    ("JobFinder — Agentic Job-Search Platform", "https://github.com/mikeyisa/jobfinder-yisa-2026",
     ["Built a multi-agent system in FastAPI and the Claude API that pulls live job "
      "postings, scores each against my résumé, rewrites the résumé for the role, and "
      "drafts recruiter outreach.",
      "Handled the agent orchestration, prompt caching to keep costs down, and a single "
      "Docker container running both the API and UI for AWS deployment."]),
    ("AI Trading Copilot — Multi-Agent Market Analysis", None,
     ["Built a multi-agent system (Next.js, FastAPI, PostgreSQL, Claude API) that pulls "
      "live market data and runs order-flow, backtesting, and prediction agents to "
      "generate directional analysis."]),
]
SKILLS = [
    "Languages: C#, Python, SQL, TypeScript/JavaScript, DAX",
    "Frameworks & APIs: .NET 6/8, Angular, React, REST APIs, FastAPI, Next.js, xUnit, "
    "LINQ / Entity Framework",
    "Cloud & DevOps: AWS (Lambda, Step Functions, S3, CloudFront, SQS, Elastic "
    "Beanstalk), Azure Data Lake, Docker, Jenkins, SonarQube, Git",
    "Data & AI: SQL Server, MySQL, PostgreSQL, Pandas, Scikit-Learn, CatBoost, "
    "Power BI, Claude API and multi-agent systems",
]


def _hyperlink(p, url, text, size=9.5):
    rid = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), LINK_BLUE); rpr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; run.append(t); link.append(run)
    p._p.append(link)


def _heading(d, text):
    p = d.add_paragraph(); p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(10.5)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    for k, v in [("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "888888")]:
        bottom.set(qn(k), v)
    pbdr.append(bottom); pPr.append(pbdr)


def _bullet(d, text):
    p = d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    p.add_run(text)


def build(tailored: dict, job: dict) -> Path:
    """tailored = {headline, summary, bullets, keywords}; job = {company, title}."""
    d = Document()
    for s in d.sections:
        s.top_margin = s.bottom_margin = Pt(40); s.left_margin = s.right_margin = Pt(54)
    st = d.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

    name = d.add_paragraph(); name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("Michael Yisa"); r.bold = True; r.font.size = Pt(19)
    sub = d.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(tailored.get("headline")
                    or "Full-Stack Software Engineer  ·  AWS Certified Solutions Architect")
    r.font.size = Pt(10.5)
    c = d.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.add_run("Austin, TX  ·  ").font.size = Pt(9.5)
    _hyperlink(c, "mailto:michaelyisa2024@gmail.com", "michaelyisa2024@gmail.com")
    c.add_run("  ·  ").font.size = Pt(9.5)
    _hyperlink(c, "https://www.linkedin.com/in/michael-yisa-382a9b249", "LinkedIn")
    c.add_run("  ·  ").font.size = Pt(9.5)
    _hyperlink(c, "https://github.com/mikeyisa", "GitHub")
    c.add_run("  ·  U.S. Citizen, Security Clearance Eligible").font.size = Pt(9.5)

    _heading(d, "Summary")
    d.add_paragraph(tailored.get("summary") or "").paragraph_format.space_after = Pt(2)

    _heading(d, "Certifications")
    _bullet(d, "AWS Certified Solutions Architect – Associate (SAA-C03), 2026")
    _bullet(d, "Google Advanced Data Analytics; Google Machine Learning; Anthropic Claude Code")

    _heading(d, "Experience")
    p = d.add_paragraph(); p.add_run("SHI International Corp — Austin, TX").bold = True
    p = d.add_paragraph(); r = p.add_run("Associate Software Engineer (Full-Stack, .NET)")
    r.bold = True; t = p.add_run("\tDec 2025 – Present"); t.italic = True; t.font.size = Pt(9.5)
    for b in (tailored.get("bullets") or []):
        _bullet(d, b)
    p = d.add_paragraph(); r = p.add_run("Jr. IT Analyst Intern"); r.bold = True
    t = p.add_run("\tMay 2025 – Dec 2025"); t.italic = True; t.font.size = Pt(9.5)
    for b in JR_BULLETS:
        _bullet(d, b)

    _heading(d, "Projects")
    for title, link, bullets in PROJECTS:
        p = d.add_paragraph(); p.add_run(title).bold = True
        if link:
            p.add_run("   ("); _hyperlink(p, link, link.replace("https://", ""), size=10); p.add_run(")")
        for b in bullets:
            _bullet(d, b)

    _heading(d, "Technical Skills")
    for s in SKILLS:
        _bullet(d, s)

    _heading(d, "Education")
    d.add_paragraph("B.S. Computer Science, Minor in Applied Mathematics — Texas State "
                    "University, Dec 2025").paragraph_format.space_after = Pt(2)
    _heading(d, "Leadership")
    _bullet(d, "Vice President, Alpha Phi Alpha Fraternity (3-year member)")
    _bullet(d, "Vice President, Texas State Data Analytics Club")

    safe = "".join(ch for ch in (job.get("company", "role")) if ch.isalnum())[:24] or "role"
    out = TAILORED_DIR / f"Michael_Yisa__{safe}_{job.get('id','')}.docx"
    d.save(out)
    return out
